"""
康复智能体 Flask 后端 v2
"""
import json
import uuid
import os
import threading
import anthropic
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_from_directory
from agent import SYSTEM_PROMPT, TOOLS, process_tool_call
import llm_providers
from llm_providers import call_llm, PROVIDERS
import database as db


# ─── 基于规则的初版方案生成器（无需 AI）────────────────────────────────────────

# 通用动作库：按部位 + 阶段分类
_EXERCISE_LIBRARY = {
    "通用": {
        "急性期": [
            {"名称": "冷敷护理", "说明": "用冰袋或冷毛巾敷于患处，每次15-20分钟，有助于消肿消炎。", "次数": "每日3-4次，每次15-20分钟", "休息": "2小时间隔", "注意": "冰袋不可直接接触皮肤，需用毛巾隔离，避免冻伤", "步骤": "1. 将冰袋用毛巾包裹；2. 轻敷于患处；3. 计时15-20分钟；4. 移除后观察皮肤状态"},
            {"名称": "肌肉静态收缩", "说明": "在不移动关节的情况下收紧肌肉，维持5秒后放松，保持肌肉张力、防止萎缩。", "次数": "10次×3组，每日2-3次", "休息": "组间休息30秒", "注意": "收缩时不要憋气，全程保持正常呼吸", "步骤": "1. 找舒适体位；2. 收紧目标肌肉维持5秒；3. 缓慢放松；4. 休息2秒后重复"},
            {"名称": "深呼吸放松", "说明": "腹式深呼吸，帮助放松全身肌肉、促进血液循环、缓解疼痛和焦虑。", "次数": "每次10个深呼吸，每日3-5次", "休息": "按需", "注意": "吸气时腹部鼓起，呼气时腹部内收，呼吸要缓慢均匀", "步骤": "1. 仰卧或坐位；2. 鼻腔缓慢吸气4秒；3. 屏息1秒；4. 嘴巴缓慢呼气6秒"},
        ],
        "亚急性期": [
            {"名称": "关节温和活动", "说明": "在无痛范围内进行关节的缓慢屈伸，逐步恢复关节活动度。", "次数": "10次×3组，每日2次", "休息": "组间休息30秒", "注意": "遇到疼痛立即停止，不强行突破疼痛范围", "步骤": "1. 取舒适体位；2. 缓慢向一个方向活动关节；3. 在无痛范围内活动；4. 缓慢返回起始位置"},
            {"名称": "肌肉激活训练", "说明": "通过轻微阻力训练激活周围肌群，为后续力量训练打基础。", "次数": "15次×3组，每日1-2次", "休息": "组间休息45秒", "注意": "动作缓慢控制，感受肌肉用力，避免借力代偿", "步骤": "1. 选择轻量阻力；2. 缓慢做目标动作；3. 在最大收缩位维持1秒；4. 缓慢还原"},
            {"名称": "静态拉伸", "说明": "对紧张肌群进行静态拉伸，改善柔韧性和血液循环。", "次数": "每侧维持30秒×3组，每日1-2次", "休息": "组间休息15秒", "注意": "拉伸感应为舒适的张力，不应有疼痛感", "步骤": "1. 缓慢进入拉伸体位；2. 维持30秒；3. 深呼吸放松；4. 缓慢还原"},
        ],
        "恢复期": [
            {"名称": "功能性力量训练", "说明": "针对受伤部位及周围肌群进行中等强度力量训练，恢复运动功能。", "次数": "12次×4组，每日1次", "休息": "组间休息60秒", "注意": "重量从轻开始，确保动作规范后再逐步增加", "步骤": "1. 热身5分钟；2. 按规范完成动作；3. 全程控制速度；4. 完成后拉伸放松"},
            {"名称": "平衡与本体感觉训练", "说明": "单腿站立或使用不稳定平面，训练平衡能力和神经肌肉控制。", "次数": "每侧30秒×3组，每日1次", "休息": "组间休息30秒", "注意": "旁边备好支撑物，初期可扶墙练习", "步骤": "1. 单腿站立；2. 保持平衡30秒；3. 逐渐加入动态挑战；4. 双侧交替进行"},
            {"名称": "有氧耐力训练", "说明": "低冲击有氧运动（游泳、骑车、快走），提升整体体能。", "次数": "20-30分钟/次，每周3-4次", "休息": "隔天进行", "注意": "心率控制在最大心率的60-70%，不要过度疲劳", "步骤": "1. 热身5分钟慢走；2. 逐渐提升强度；3. 保持目标心率；4. 放松5分钟"},
        ],
        "强化期": [
            {"名称": "专项力量强化", "说明": "针对运动需求的专项力量训练，恢复竞技状态。", "次数": "8-10次×4-5组，每日1次", "休息": "组间休息90秒", "注意": "重量递增原则，技术动作优先于重量", "步骤": "1. 充分热身；2. 中大重量规范训练；3. 注意保护措施；4. 完成后充分放松"},
            {"名称": "爆发力训练", "说明": "增强式训练，恢复快速力量爆发能力，为重返运动场做准备。", "次数": "6-8次×3-4组，每周2-3次", "休息": "组间休息2分钟", "注意": "疲劳状态下不做爆发力训练，注意落地动作规范", "步骤": "1. 充分热身；2. 以最快速度完成动作；3. 完全休息后再做下一组；4. 关注动作质量"},
            {"名称": "专项运动恢复训练", "说明": "模拟实际运动场景的专项动作练习，测试功能恢复情况。", "次数": "视项目而定，每日1次", "休息": "按需休息", "注意": "循序渐进，确认无痛后才进行专项训练", "步骤": "1. 从低速低强度开始；2. 逐步加快速度；3. 最后进行对抗性练习；4. 记录任何不适"},
        ],
    },
    "膝关节": {
        "急性期": [
            {"名称": "股四头肌等长收缩", "说明": "平躺，腿伸直，大腿肌肉用力向下压，维持5秒后放松。有效防止肌肉萎缩。", "次数": "10次×3组，每日3次", "休息": "组间休息20秒", "注意": "过程中膝关节不要弯曲，感受大腿前侧肌肉收紧", "步骤": "1. 平躺伸腿；2. 膝盖用力向床面下压；3. 维持5秒；4. 放松2秒重复"},
        ],
        "亚急性期": [
            {"名称": "直腿抬高", "说明": "仰卧，一腿屈曲踩地，另一腿伸直抬高至45度，维持2秒后缓慢放下。", "次数": "15次×3组，每日2次", "休息": "组间休息30秒", "注意": "抬腿时腰部不要离地，腹肌保持收紧", "步骤": "1. 仰卧屈膝；2. 患腿伸直；3. 缓慢抬高45度；4. 维持2秒后缓慢放下"},
            {"名称": "迷你深蹲", "说明": "双脚与肩同宽，屈膝下蹲至15-30度，膝关节不超过脚尖，起身还原。", "次数": "10次×3组，每日1次", "休息": "组间休息45秒", "注意": "下蹲深度以无痛为准，膝盖保持与脚尖同方向", "步骤": "1. 双脚与肩同宽站立；2. 下蹲时重心均匀；3. 膝盖对准小脚趾方向；4. 起身时臀部夹紧"},
        ],
        "恢复期": [
            {"名称": "标准深蹲", "说明": "逐步增加深蹲深度至90度，强化股四头肌和臀部肌群。", "次数": "12次×4组，每日1次", "休息": "组间休息60秒", "注意": "全脚掌踩地，膝盖不内扣，腰背保持中立位", "步骤": "1. 站位准备；2. 缓慢下蹲至90度；3. 维持1秒；4. 有力起身"},
            {"名称": "踏台阶训练", "说明": "单腿踩上台阶，另一腿随后抬起，锻炼平衡和腿部力量。", "次数": "每侧12次×3组，每日1次", "休息": "组间休息45秒", "注意": "台阶高度从低开始，保持身体垂直，不要向前倾斜", "步骤": "1. 站于台阶前；2. 患腿单腿踩上；3. 身体重心转移；4. 缓慢下台"},
        ],
        "强化期": [
            {"名称": "单腿深蹲", "说明": "单腿站立缓慢下蹲，测试和强化膝关节稳定性和力量。", "次数": "每侧8次×3组，每日1次", "休息": "组间休息60秒", "注意": "必须在膝关节完全稳定后才进行，旁边保持支撑物", "步骤": "1. 单腿站立；2. 缓慢下蹲至45-60度；3. 膝关节稳定不晃动；4. 有力起身"},
        ],
    },
    "肩关节": {
        "急性期": [
            {"名称": "钟摆运动", "说明": "上身前倾，手臂放松自然下垂，靠身体轻微晃动带动手臂做钟摆样运动，减轻肩关节压力。", "次数": "每方向20次×3组，每日2-3次", "休息": "组间休息20秒", "注意": "手臂完全放松，不主动用力，靠重力自然摆动", "步骤": "1. 健手扶桌前倾；2. 患臂自然下垂；3. 身体轻微摆动带动手臂；4. 前后、左右、画圈各20次"},
        ],
        "亚急性期": [
            {"名称": "爬墙练习", "说明": "面向墙壁站立，手指沿墙面缓慢向上爬行，逐步提升肩关节上举范围。", "次数": "10次×3组，每日2次", "休息": "组间休息30秒", "注意": "到达疼痛点时停止，不强行拉高，每天记录最高点", "步骤": "1. 面墙站立距30cm；2. 手指从腰部开始；3. 缓慢向上爬行；4. 到最高无痛点停留3秒后还原"},
            {"名称": "肩胛骨收缩", "说明": "坐直，两侧肩胛骨向中间夹紧，维持5秒，改善肩关节稳定性。", "次数": "15次×3组，每日2次", "休息": "组间休息20秒", "注意": "动作时不要耸肩，保持肩膀自然放松下沉", "步骤": "1. 坐直双手自然放腿上；2. 肩胛向脊柱方向夹紧；3. 维持5秒；4. 缓慢放松"},
        ],
        "恢复期": [
            {"名称": "肩部环绕", "说明": "手臂大幅度画圆活动肩关节，全方向恢复活动度。", "次数": "每方向10次×3组，每日1次", "休息": "组间休息30秒", "注意": "范围从小到大，在无痛范围内活动", "步骤": "1. 站立手臂放松；2. 向前大幅画圆10次；3. 向后大幅画圆10次；4. 感受全程无痛"},
            {"名称": "弹力带外旋训练", "说明": "使用轻阻力弹力带做肩关节外旋动作，强化肩袖肌群。", "次数": "15次×3组，每日1次", "休息": "组间休息45秒", "注意": "肘部夹紧身体，只做外旋动作，不要耸肩", "步骤": "1. 肘弯90度夹腰；2. 手持弹力带；3. 向外旋转至极限；4. 缓慢还原"},
        ],
        "强化期": [
            {"名称": "哑铃肩部推举", "说明": "端坐或站立，哑铃从耳旁推举至头顶，全面强化肩部力量。", "次数": "10次×4组，每日1次", "休息": "组间休息90秒", "注意": "重量从轻开始，动作轨迹规范，不要弓腰借力", "步骤": "1. 坐直核心收紧；2. 哑铃在耳旁；3. 垂直推举过头；4. 缓慢还原"},
        ],
    },
    "腰部": {
        "急性期": [
            {"名称": "卧床休息 + 体位调整", "说明": "急性期以卧床休息为主，侧卧时双膝间夹枕头，减轻腰椎压力。", "次数": "全天姿势管理", "休息": "按需", "注意": "避免久坐，每30分钟改变体位，不要弯腰搬重物", "步骤": "1. 侧卧双膝微弯；2. 膝间夹枕头；3. 上方手臂放体前；4. 每30分钟翻身一次"},
            {"名称": "骨盆钟摆", "说明": "仰卧屈膝，骨盆前倾后倾交替活动，放松腰椎周围肌肉。", "次数": "10次×3组，每日2-3次", "休息": "组间休息20秒", "注意": "动作要小而缓慢，疼痛加重立即停止", "步骤": "1. 仰卧屈膝；2. 腰部慢慢向床面压；3. 再慢慢弓起腰；4. 交替进行"},
        ],
        "亚急性期": [
            {"名称": "猫牛式", "说明": "四肢支撑，交替做脊柱屈伸（猫式：拱背，牛式：塌腰），改善腰椎活动度。", "次数": "10次×3组，每日2次", "休息": "组间休息20秒", "注意": "动作要缓慢，配合呼吸，感受脊柱逐节活动", "步骤": "1. 四点支撑；2. 拱背向上（呼气）；3. 塌腰向下（吸气）；4. 缓慢交替"},
            {"名称": "鸟狗式", "说明": "四点支撑，对侧手腿同时缓慢伸展，增强核心稳定性。", "次数": "每侧10次×3组，每日1次", "休息": "组间休息30秒", "注意": "骨盆保持水平，不要转动，腰部不能塌陷", "步骤": "1. 四点支撑；2. 右手左腿同时伸直；3. 维持3秒；4. 还原后换另一侧"},
        ],
        "恢复期": [
            {"名称": "平板支撑", "说明": "前臂支撑，保持身体一条直线，强化核心肌群。", "次数": "30-60秒×3组，每日1次", "休息": "组间休息60秒", "注意": "腰部不要下沉，也不要抬高臀部，保持正常呼吸", "步骤": "1. 前臂撑地；2. 脚尖点地身体成直线；3. 收腹收臀；4. 保持稳定呼吸"},
            {"名称": "臀桥", "说明": "仰卧屈膝，臀部发力向上顶，强化臀部和腰椎稳定肌群。", "次数": "15次×3组，每日1次", "休息": "组间休息45秒", "注意": "靠臀部用力而非腰部，到达最高点时夹紧臀部", "步骤": "1. 仰卧屈膝；2. 脚踩地臀部发力向上；3. 最高点夹紧臀部；4. 缓慢下降"},
        ],
        "强化期": [
            {"名称": "硬拉强化", "说明": "从地面拾起重物的动作模式训练，强化腰部和全链条力量。", "次数": "8次×4组，每日1次", "休息": "组间休息90秒", "注意": "腰背始终保持中立位，不要弯腰圆背，重量循序渐进", "步骤": "1. 站位屈髋下蹲；2. 保持腰背中立；3. 臀腿发力站起；4. 缓慢还原"},
        ],
    },
    "踝关节": {
        "急性期": [
            {"名称": "踝泵练习", "说明": "卧位，踝关节做上下泵压动作，促进血液循环，防止水肿。", "次数": "每分钟20次，每次5分钟，每日多次", "休息": "按需", "注意": "以疼痛为界，不要勉强活动范围，可同时热敷辅助", "步骤": "1. 平躺或坐位；2. 脚背向上勾起；3. 缓慢下压绷脚；4. 持续有节律进行"},
        ],
        "亚急性期": [
            {"名称": "踝关节画字母", "说明": "坐位，脚悬空，用大拇指在空中画英文字母A-Z，恢复踝关节全方向活动度。", "次数": "A-Z全套×3，每日2次", "休息": "组间休息20秒", "注意": "动作来自踝关节，不要用腿部代偿", "步骤": "1. 坐位腿伸出；2. 以踝为轴；3. 大拇指在空中写字母；4. 尽量使关节活动充分"},
            {"名称": "提踵训练", "说明": "双手扶墙，脚跟抬起用前脚掌支撑，训练小腿肌肉力量。", "次数": "15次×3组，每日1-2次", "休息": "组间休息30秒", "注意": "缓慢上升下降，不要跳跃式落下，感受小腿肌肉发力", "步骤": "1. 双手扶墙；2. 脚跟缓慢抬起；3. 到最高位维持1秒；4. 缓慢放下"},
        ],
        "恢复期": [
            {"名称": "单腿平衡练习", "说明": "单脚站立，训练踝关节本体感觉和稳定性。", "次数": "每侧30秒×3组，每日1次", "休息": "组间休息20秒", "注意": "初期可扶墙辅助，进步后加入闭眼挑战", "步骤": "1. 单脚站立；2. 保持平衡30秒；3. 进阶可轻微晃动；4. 双侧交替"},
        ],
        "强化期": [
            {"名称": "跳绳 / 弹跳训练", "说明": "从双脚低强度跳跃开始，逐步恢复踝关节承重和弹跳能力。", "次数": "每次3-5分钟，每周3次", "休息": "组间充分休息", "注意": "确保踝关节完全无痛后才开始跳跃，软地面开始", "步骤": "1. 双脚小幅跳跃热身；2. 逐步增加幅度；3. 单脚跳进阶；4. 确认无痛无不适"},
        ],
    },
}

# 关键词到部位的映射
_KEYWORD_TO_REGION = {
    "膝": "膝关节", "膝盖": "膝关节", "半月板": "膝关节", "韧带": "膝关节",
    "前交叉": "膝关节", "后交叉": "膝关节", "髌骨": "膝关节",
    "肩": "肩关节", "肩袖": "肩关节", "肩膀": "肩关节", "肩锁": "肩关节",
    "腰": "腰部", "腰椎": "腰部", "腰肌": "腰部", "椎间盘": "腰部", "脊柱": "腰部",
    "踝": "踝关节", "脚踝": "踝关节", "踝关节": "踝关节",
}

def generate_default_plan(patient_data: dict) -> dict:
    """根据患者档案生成规则驱动的初版康复方案（无需 AI）"""
    name = patient_data.get("name", "")
    injury = patient_data.get("injury_description", "")
    stage = patient_data.get("recovery_stage", "急性期")
    pain_level = patient_data.get("pain_level") or 5

    # 识别受伤部位
    region = "通用"
    for kw, reg in _KEYWORD_TO_REGION.items():
        if kw in injury:
            region = reg
            break

    # 确保阶段有效
    valid_stages = ["急性期", "亚急性期", "恢复期", "强化期"]
    if stage not in valid_stages:
        stage = "急性期"

    # 获取动作：优先部位专项，补充通用动作
    specific = _EXERCISE_LIBRARY.get(region, {}).get(stage, [])
    generic = _EXERCISE_LIBRARY["通用"].get(stage, [])
    # 合并：最多取6个动作，优先专项
    actions = (specific + generic)[:6]

    # 疼痛高时减少动作量
    if pain_level >= 7:
        actions = actions[:3]
        intensity_note = "当前疼痛评分较高，初期动作以轻度为主，以控制疼痛为第一目标。"
    elif pain_level >= 4:
        actions = actions[:5]
        intensity_note = "训练过程中以微痛为界，不强行突破，循序渐进。"
    else:
        intensity_note = "疼痛控制良好，可按计划推进，注意动作质量。"

    # 各阶段目标
    goals = {
        "急性期":   f"控制炎症和疼痛，防止肌肉萎缩，维持周围关节活动度。{intensity_note}",
        "亚急性期": f"逐步恢复关节活动范围，激活周围肌群，开始轻度功能训练。{intensity_note}",
        "恢复期":   f"恢复正常关节活动度和基础肌肉力量，提升功能性运动能力。{intensity_note}",
        "强化期":   f"全面恢复运动能力，针对专项需求进行力量和爆发力训练，准备重返运动场。{intensity_note}",
    }

    principles = {
        "急性期":   "POLICE 原则：保护（Protection）→ 适当负重（Optimal Loading）→ 冰敷（Ice）→ 加压（Compression）→ 抬高（Elevation）",
        "亚急性期": "无痛原则：所有动作在无痛或微痛范围内进行，不强行拉伸，动作缓慢控制",
        "恢复期":   "渐进超负荷：逐步增加训练量和强度，注重动作质量，双侧力量均衡",
        "强化期":   "功能专项：模拟真实运动需求，兼顾力量、速度、协调性，做好预防再伤措施",
    }

    contraindications = {
        "急性期":   ["禁止热敷（24小时内）", "避免任何加剧疼痛的动作", "不要强行活动受伤关节", "严禁带痛运动"],
        "亚急性期": ["不强行突破疼痛范围", "避免高冲击运动（跑步、跳跃）", "禁止患处直接受压过久"],
        "恢复期":   ["避免过度训练，注意疲劳信号", "出现肿胀及时冰敷并减量", "不跳跃性增加重量"],
        "强化期":   ["避免疲劳时进行爆发力训练", "出现原受伤处不适立即停训咨询", "做好热身和拉伸"],
    }

    region_display = region if region != "通用" else ("根据伤情" if injury else "全身")

    plan = {
        "injury_type": injury or "运动训练伤",
        "stage": stage,
        "severity": f"疼痛评分 {pain_level}/10",
        "region": region_display,
        "目标": goals.get(stage, ""),
        "原则": principles.get(stage, ""),
        "动作": actions,
        "禁忌": contraindications.get(stage, []),
        "预期时间": {"急性期": "1-2周", "亚急性期": "2-4周", "恢复期": "4-8周", "强化期": "4-12周"}.get(stage, "视康复进展"),
        "说明": f"本方案为根据【{name}】档案信息自动生成的初版方案（部位：{region_display}，阶段：{stage}）。建议完成后与康复师或智能助手对话，进一步细化和个性化调整。",
        "is_auto_generated": True,
    }
    return plan


def _call_ai_for_plan(patient_data: dict, base_plan: dict) -> dict:
    """调用 AI 对 base_plan 进行个性化增强，失败时原样返回 base_plan。"""
    cfg = db.get_system_config()
    api_key = cfg.get("ai_api_key", "").strip()
    if not api_key:
        return base_plan

    provider = cfg.get("ai_provider", "claude")
    model    = cfg.get("ai_model", "claude-sonnet-4-6")
    base_url = cfg.get("ai_base_url", "").strip() or None

    name     = patient_data.get("name", "患者")
    injury   = patient_data.get("injury_description", "")
    stage    = patient_data.get("recovery_stage", "恢复期")
    pain     = patient_data.get("pain_level") or 3
    age      = patient_data.get("age") or 20
    gender   = patient_data.get("gender", "")
    med_hist = patient_data.get("medical_history", "") or "无"
    surg     = patient_data.get("surgery_history", "") or "无"
    unit     = patient_data.get("unit", "")

    prompt = f"""你是一名专业运动医学康复专家，请根据患者档案，为以下四阶段康复计划生成个性化内容。

## 患者信息
- 姓名：{name}，性别：{gender}，年龄：{age}岁
- 伤情描述：{injury}
- 当前康复阶段：{stage}，疼痛评分：{pain}/10
- 既往病史：{med_hist}，手术史：{surg}
- 所在单位：{unit or "未填写"}

## 输出格式（严格按照 JSON 返回，不要任何额外文字）
{{
  "overall_principles": ["原则1","原则2","原则3","原则4","原则5"],
  "phases": {{
    "急性期":  {{"goal":"...","rationale":"...（结合患者具体伤情，3-5句）","exercises":[{{"name":"...","freq":"...","focus":"..."}}],"key_points":["...","...","...","..."]}},
    "亚急性期":{{"goal":"...","rationale":"...","exercises":[...],"key_points":[...]}},
    "恢复期":  {{"goal":"...","rationale":"...","exercises":[...],"key_points":[...]}},
    "强化期":  {{"goal":"...","rationale":"...","exercises":[...],"key_points":[...]}}
  }},
  "return_criteria": ["标准1","标准2","标准3","标准4","标准5"]
}}

每个阶段 exercises 提供 4-5 个动作，key_points 提供 4 条注意事项。"""

    sys_msg = "你是专业康复医学专家，只输出严格的 JSON 数据，不输出任何解释性文字。"
    messages = [{"role": "system", "content": sys_msg},
                {"role": "user",   "content": prompt}]

    try:
        if base_url or provider != "claude":
            b_url = base_url or PROVIDERS.get(provider, {}).get("base_url", "https://api.openai.com/v1")
            resp_data = llm_providers._openai_chat(b_url, api_key, model, messages, timeout=90)
            ai_text = resp_data["choices"][0]["message"].get("content", "")
        else:
            client = anthropic.Anthropic(api_key=api_key)
            resp = client.messages.create(
                model=model, max_tokens=4096,
                system=sys_msg,
                messages=[{"role": "user", "content": prompt}],
            )
            ai_text = resp.content[0].text if resp.content else ""

        ai_text = ai_text.strip()
        if "```json" in ai_text:
            ai_text = ai_text.split("```json")[1].split("```")[0].strip()
        elif "```" in ai_text:
            ai_text = ai_text.split("```")[1].split("```")[0].strip()

        ai_data = json.loads(ai_text)

        if ai_data.get("overall_principles"):
            base_plan["overall_principles"] = ai_data["overall_principles"]
        if ai_data.get("return_criteria"):
            base_plan["return_criteria"]["items"] = ai_data["return_criteria"]

        ai_phases = ai_data.get("phases", {})
        for phase in base_plan.get("phases", []):
            s = phase["stage"]
            if s in ai_phases:
                ap = ai_phases[s]
                for key in ("goal", "rationale", "exercises", "key_points"):
                    if ap.get(key):
                        phase[key] = ap[key]

        base_plan["ai_generated"] = True
        base_plan["is_auto_generated"] = True
        base_plan["generated_at"] = datetime.now().strftime("%Y年%m月%d日 %H:%M")
        return base_plan

    except Exception as e:
        print(f"[warn] AI 计划增强失败，使用规则方案: {e}")
        return base_plan


def _generate_and_save_ai_plan(patient_data: dict, patient_id: int, sid):
    """后台线程：生成 AI 方案并覆盖保存。"""
    session_state["ai_plan_status"] = "generating"
    try:
        base_plan = generate_full_training_plan(patient_data)
        final_plan = _call_ai_for_plan(patient_data, base_plan)
        db.save_rehab_plan(sid, patient_id, final_plan)
        session_state["ai_plan_status"] = "done"
        print(f"[info] AI 训练计划已更新 patient_id={patient_id}")
    except Exception as e:
        session_state["ai_plan_status"] = "idle"
        print(f"[warn] 后台 AI 计划线程失败: {e}")


app = Flask(__name__)
app.secret_key = "rehab_v2_secret_key_2024"
app.json.sort_keys = False   # 保持 JSON 键顺序（如衣食行顺序）

# 全局会话状态（单用户 Demo）
session_state = {
    "patient": None,
    "messages": [],
    "rehab_plan": None,
    "feedback_log": [],
    "session_id": None,
    "patient_id": None,
    "ai_plan_status": "idle",   # idle | generating | done
}

# 启动时自动初始化数据库并绑定演示患者
try:
    db.init_db()
    _seed_pid = db.ensure_seed_patient()
    if _seed_pid and not session_state["patient_id"]:
        session_state["patient_id"] = _seed_pid
        _p = db.get_patient(_seed_pid)
        if _p:
            session_state["patient"] = dict(_p)
except Exception as _e:
    print(f"[warn] 自动初始化失败: {_e}")


def _ensure_session(provider: str = "claude", model: str = "claude-sonnet-4-6"):
    if not session_state["session_id"]:
        sid = str(uuid.uuid4())
        session_state["session_id"] = sid
        db.create_session(sid, provider, model)
    return session_state["session_id"]


def run_agent_turn(user_message: str, api_key: str,
                   provider: str = "claude", model: str = "claude-sonnet-4-6") -> dict:
    session_id = _ensure_session(provider, model)

    session_state["messages"].append({
        "role": "user",
        "content": user_message
    })

    def process_tool(name, inp):
        return process_tool_call(name, inp, session_state)

    final_text, tool_calls_made = call_llm(
        provider, model, api_key,
        session_state["messages"],
        TOOLS,
        SYSTEM_PROMPT,
        process_tool,
    )

    patient_id = session_state.get("patient_id")
    for call in tool_calls_made:
        tool_name = call.get("tool", "")
        tool_inp = call.get("input", {})
        tool_res = call.get("result", {})

        if tool_name == "save_patient_info":
            pid = db.upsert_patient_from_tool(session_id, tool_inp)
            session_state["patient_id"] = pid
            patient_id = pid
        elif tool_name == "generate_rehab_plan":
            plan = session_state.get("rehab_plan") or tool_res.get("plan", {})
            db.save_rehab_plan(session_id, patient_id, plan)
        elif tool_name == "log_training_feedback":
            db.log_feedback(session_id, patient_id, tool_inp)

        db.log_tool_call(session_id, patient_id, tool_name, tool_inp, tool_res)

    return {
        "response": final_text,
        "tool_calls": tool_calls_made,
        "patient": session_state["patient"],
        "rehab_plan": session_state["rehab_plan"]
    }


# ─── 页面路由 ─────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/profile")
def profile():
    return render_template("profile.html")


@app.route("/checkin")
def checkin():
    return render_template("checkin.html")


@app.route("/plan")
def plan():
    return render_template("plan.html")


@app.route("/knowledge")
def knowledge():
    return render_template("knowledge.html")


@app.route("/companion")
def companion():
    return render_template("companion.html")


@app.route("/teacher")
def teacher():
    return render_template("teacher.html")


@app.route("/admin")
def admin():
    return render_template("admin.html")


# ─── 患者档案 API ─────────────────────────────────────────────────────────────

@app.route("/api/plan/ai-status", methods=["GET"])
def get_ai_plan_status():
    status = session_state.get("ai_plan_status", "idle")
    if status == "done":
        session_state["ai_plan_status"] = "idle"   # 消费后重置
    return jsonify({"status": status})


@app.route("/api/me", methods=["GET"])
def get_me():
    patient_id = session_state.get("patient_id")
    if patient_id:
        patient = db.get_patient(patient_id)
        if patient:
            return jsonify(patient)
    if session_state.get("patient"):
        return jsonify(session_state["patient"])
    return jsonify(None)


@app.route("/api/me", methods=["POST"])
def update_me():
    data = request.get_json() or {}
    if not data.get("name"):
        return jsonify({"error": "姓名不能为空"}), 400

    patient_id = session_state.get("patient_id")
    is_new = not patient_id
    if patient_id:
        db.update_patient(patient_id, data)
    else:
        patient_id = db.create_patient(data)
        session_state["patient_id"] = patient_id
        sid = session_state.get("session_id")
        if sid:
            db.update_session_patient(sid, patient_id, data.get("name", ""))

    session_state["patient"] = {
        "name": data.get("name", ""),
        "gender": data.get("gender", ""),
        "age": data.get("age"),
        "height": data.get("height"),
        "weight": data.get("weight"),
        "injury_description": data.get("injury_description", ""),
        "recovery_stage": data.get("recovery_stage", "未知"),
        "pain_level": data.get("pain_level"),
        "notes": data.get("notes", ""),
    }

    # 有伤情信息时自动生成/更新训练方案（每次保存档案都重新生成）
    plan_generated = False
    if data.get("injury_description") or data.get("recovery_stage"):
        sid = session_state.get("session_id") or None
        # 先存规则方案，保证响应速度
        rule_plan = generate_full_training_plan(data)
        db.save_rehab_plan(sid, patient_id, rule_plan)
        plan_generated = True
        # 后台线程用 AI 覆盖更新
        cfg = db.get_system_config()
        if cfg.get("ai_api_key", "").strip():
            t = threading.Thread(
                target=_generate_and_save_ai_plan,
                args=(data, patient_id, sid),
                daemon=True,
            )
            t.start()

    return jsonify({
        "success": True,
        "patient_id": patient_id,
        "message": "档案已更新",
        "plan_generated": plan_generated,
        "ai_plan_generating": plan_generated and bool(db.get_system_config().get("ai_api_key", "").strip()),
    })


# ─── 详细训练计划生成器 ────────────────────────────────────────────────────────

def generate_full_training_plan(patient: dict) -> dict:
    """
    根据患者档案生成完整的分周期训练计划，
    包含：各阶段内容、训练理由、注意事项、衣食行三方面提示。
    """
    name        = patient.get("name", "")
    injury      = patient.get("injury_description", "")
    stage       = patient.get("recovery_stage", "恢复期")
    pain        = patient.get("pain_level") or 3
    age         = patient.get("age") or 20
    gender      = patient.get("gender", "")
    unit        = patient.get("unit", "")
    med_history = patient.get("medical_history", "")

    # 识别受伤部位
    region = "通用"
    for kw, reg in _KEYWORD_TO_REGION.items():
        if kw in injury:
            region = reg
            break

    # 判断受伤严重程度
    severity_level = "轻度" if pain <= 3 else "中度" if pain <= 6 else "重度"

    # 生成各阶段计划（从当前阶段开始往后）
    stage_order = ["急性期", "亚急性期", "恢复期", "强化期"]
    start_idx   = stage_order.index(stage) if stage in stage_order else 2

    # 各阶段的详细描述（通用底板 + 部位专项叠加）
    stage_details = {
        "急性期": {
            "week_range": "第1-2周",
            "theme": "炎症控制·保护制动",
            "color": "red",
            "goal": "控制局部炎症与肿胀，防止二次损伤，维持非受伤部位肌肉张力。",
            "rationale": (
                f"受伤初期（{injury[:20]}），局部组织处于炎症反应阶段，"
                "血管扩张、白细胞聚集清除坏死组织。此阶段主要任务是「保护」：用POLICE原则"
                "（保护→适当负重→冰敷→加压→抬高）控制水肿；同时通过等长收缩维持肌肉活性，"
                "为后续恢复打好基础。过度活动会加重出血和炎症，过度制动则会引发肌肉萎缩和关节粘连。"
            ),
            "exercises": [
                {"name": "冰敷护理（RICE原则）", "freq": "每日3-4次，每次15-20分钟", "focus": "消炎消肿"},
                {"name": "等长肌肉收缩", "freq": "10次×3组，每日2-3次", "focus": "防止萎缩"},
                {"name": "非受伤关节活动", "freq": "按需进行", "focus": "保持整体循环"},
                {"name": "抬高患肢休息", "freq": "休息时执行", "focus": "促进回流减肿"},
            ],
            "key_points": [
                "24小时内禁止热敷，会加重出血和肿胀",
                "冰袋需包裹毛巾，每次不超过20分钟，防止冻伤",
                "避免任何引起疼痛加重的动作",
                "如出现持续剧烈疼痛、麻木或皮肤变色，立即就医",
            ],
        },
        "亚急性期": {
            "week_range": "第3-4周",
            "theme": "组织修复·活动恢复",
            "color": "orange",
            "goal": "逐步恢复关节全范围活动度，激活周围肌群，开始低强度功能训练。",
            "rationale": (
                "伤后1-2周，炎症逐渐消退，成纤维细胞开始修复受损组织，胶原纤维无序排列、质量较脆弱。"
                "此阶段需要「适当应力刺激」：轻度的拉伸和运动可引导胶原纤维沿受力方向排列，"
                "提高修复质量；同时通过关节活动度练习和肌肉激活训练，避免关节僵硬和萎缩进一步加重。"
                "强度不足会延误恢复，强度过大会撕裂脆弱的修复组织。"
            ),
            "exercises": [
                {"name": "关节温和主动活动", "freq": "10次×3组，每日2次", "focus": "恢复活动度"},
                {"name": "肌肉激活训练", "freq": "15次×3组，每日1-2次", "focus": "激活萎缩肌群"},
                {"name": "静态拉伸", "freq": "每侧30秒×3组，每日1-2次", "focus": "改善柔韧性"},
                {"name": "步行训练（平地）", "freq": "每次10-15分钟，每日1-2次", "focus": "恢复步态"},
            ],
            "key_points": [
                "活动范围以无痛为界，切勿强行突破",
                "训练后若出现肿胀加重，降低下次训练量并冰敷",
                "禁止跑步、跳跃等高冲击活动",
                "每日记录疼痛评分，如持续≥6分超过3天需暂停并就医",
            ],
        },
        "恢复期": {
            "week_range": "第5-8周",
            "theme": "力量重建·功能恢复",
            "color": "blue",
            "goal": "恢复受伤部位力量至对侧80%以上，建立稳定的运动控制能力，恢复日常功能性动作。",
            "rationale": (
                "此阶段胶原组织修复接近完成，虽然强度尚未完全恢复，但已能承受渐进性力量训练。"
                f"针对{region}损伤，本体感觉（本体感受器）功能受损是复发的主要原因，"
                "因此本阶段重点在于「力量+感觉双重重建」：通过渐进阻力训练恢复肌肉力量，"
                "通过平衡和本体感觉训练重建神经肌肉控制，二者缺一不可。"
                f"当前疼痛评分{pain}/10，提示组织恢复状态良好，可按计划推进。"
            ),
            "exercises": [
                {"name": "深蹲 / 分腿蹲", "freq": "12次×4组，每日1次", "focus": "恢复腿部综合力量"},
                {"name": "踏台阶训练", "freq": "每侧12次×3组，每日1次", "focus": "单腿力量与平衡"},
                {"name": "弹力带抗阻训练", "freq": "15次×3组，每日1次", "focus": "专项肌群强化"},
                {"name": "平衡板 / 单腿站立", "freq": "每侧30秒×3组，每日1次", "focus": "本体感觉重建"},
                {"name": "快步走 / 椭圆机", "freq": "每次20-30分钟，每周4次", "focus": "有氧耐力基础"},
            ],
            "key_points": [
                "力量训练遵循渐进超负荷：先保证动作规范，再逐步增加重量或难度",
                "训练后冰敷10分钟，有助于控制微小炎症反应",
                "双侧力量差异超过15%前，不得开始跑步训练",
                "出现关节积液（肿胀明显）立即降量，持续不消退需就医",
            ],
        },
        "强化期": {
            "week_range": "第9-12周",
            "theme": "专项强化·重返训练",
            "color": "green",
            "goal": f"全面恢复运动能力，通过专项测试后重返{'部队训练' if unit else '正常训练'}，建立长效预防机制。",
            "rationale": (
                "恢复期结束后，组织强度已接近正常水平，但运动模式和神经控制仍需专项训练巩固。"
                f"针对部队训练需求，本阶段以「专项力量+爆发力+重返测试」为核心：通过大重量力量训练、"
                "增强式训练和多方向变向练习，模拟实战训练强度；通过标准化功能测试（如单腿跳距比、"
                "Y-balance测试）客观评估恢复程度，确保以安全的身体状态重返训练场，避免二次受伤。"
            ),
            "exercises": [
                {"name": "单腿深蹲强化", "freq": "每侧8次×4组，每日1次", "focus": "终极力量检验"},
                {"name": "慢跑训练（渐进距离）", "freq": "从500米起，每周加200米", "focus": "跑步功能恢复"},
                {"name": "侧向变向训练", "freq": "6次×3组，每周2-3次", "focus": "运动专项模式"},
                {"name": "爆发力训练（跳跃）", "freq": "6-8次×3组，每周2次", "focus": "快速力量恢复"},
                {"name": "专项预防性训练", "freq": "每日热身10分钟", "focus": "长效预防再伤"},
            ],
            "key_points": [
                "慢跑训练首次进行时需有人陪同，出现不适立即停止",
                "变向和爆发力训练必须在无痛状态下进行，不带伤训练",
                "重返训练前完成功能测试：单腿跳距双侧差异<10%方可通过",
                "此后长期保持每日预防性热身，这是防止复发的关键",
            ],
        },
    }

    # 构建完整四阶段时间线（所有阶段，标注当前/已完成/待进行）
    week_spans = {0: (1, 2), 1: (3, 4), 2: (5, 8), 3: (9, 12)}
    phases = []
    for i, s in enumerate(stage_order):
        detail = stage_details[s]
        w_start, w_end = week_spans[i]
        if i < start_idx:
            status = "done"
            week_label = f"第{w_start}-{w_end}周（已完成）"
        elif i == start_idx:
            status = "current"
            week_label = f"第{w_start}-{w_end}周（当前阶段）"
        else:
            status = "upcoming"
            week_label = f"第{w_start}-{w_end}周（待进行）"

        phases.append({
            "stage": s,
            "week_label": week_label,
            "status": status,
            "is_current": (status == "current"),
            "theme": detail["theme"],
            "color": detail["color"],
            "goal": detail["goal"],
            "rationale": detail["rationale"],
            "exercises": detail["exercises"],
            "key_points": detail["key_points"],
        })

    # 总预计恢复时间
    remaining_stages = len(stage_order) - start_idx
    total_weeks_map = {0: "10-14周", 1: "8-10周", 2: "6-8周", 3: "4-6周"}
    total_weeks = total_weeks_map.get(start_idx, "6-8周")

    # 衣食行三方面注意事项（按部位定制）
    lifestyle_tips = _get_lifestyle_tips(region, stage, pain, age, gender, unit)

    return {
        "patient_name": name,
        "injury": injury,
        "region": region,
        "current_stage": stage,
        "severity": severity_level,
        "pain": pain,
        "total_weeks": total_weeks,
        "phases": phases,
        "lifestyle": lifestyle_tips,
        "overall_principles": [
            "循序渐进：每个阶段训练量不超过上一阶段的10-15%增幅",
            "无痛优先：任何训练以疼痛≤3分为界，超过立即调整",
            "双侧均衡：注意健患侧力量对比，差距缩小至15%内才进阶",
            "记录追踪：每日记录训练完成情况和疼痛变化，供复核参考",
            "主动沟通：出现异常症状（持续肿胀、麻木、疼痛加重）立即上报医官",
        ],
        "return_criteria": {
            "title": "重返训练达标标准",
            "items": [
                "患侧肌力≥健侧90%（等速肌力测试）",
                "单腿跳距比（患/健）≥90%",
                "Y-balance测试综合评分≥94%",
                "连续跑步1000米无疼痛或肿胀",
                "完成完整连队训练科目无不适",
            ]
        },
        "generated_at": datetime.now().strftime("%Y年%m月%d日"),
    }


def _get_lifestyle_tips(region, stage, pain, age, gender, unit):
    """生成衣食行三方面生活注意事项"""
    from datetime import datetime

    # 膝关节专项提示（兼顾通用）
    knee_tips = {
        "衣": {
            "icon": "👕",
            "title": "服装与防护装备",
            "tips": [
                {"label": "护膝使用", "content": "训练时佩戴弹性护膝（压缩型），提供本体感觉支持和轻度保护。不建议24小时佩戴，影响血液循环。", "level": "重要"},
                {"label": "鞋具选择", "content": "选择有足弓支撑和缓震功能的运动鞋，避免平底鞋训练。建议每300公里或6个月更换一次跑鞋。", "level": "重要"},
                {"label": "训练服装", "content": "选择弹性好、透气性强的运动裤，避免过紧影响膝关节活动范围。寒冷天气加穿保暖护膝。", "level": "一般"},
                {"label": "保暖防寒", "content": "气温低于10°C时，训练前充分热身，佩戴保暖护膝。膝关节受寒会加重僵硬感。", "level": "一般"},
            ]
        },
        "食": {
            "icon": "🥗",
            "title": "营养与饮食建议",
            "tips": [
                {"label": "蛋白质摄入", "content": f"每日蛋白质摄入量建议1.6-2.0g/kg体重（约{round(72.5*1.8)}g/天），优先选择鸡肉、鱼肉、鸡蛋和豆制品，支持韧带修复。", "level": "重要"},
                {"label": "抗炎营养", "content": "增加Omega-3摄入：每日1-2份深海鱼（三文鱼、鲭鱼）或亚麻籽油，可有效降低局部炎症反应。", "level": "重要"},
                {"label": "维生素C", "content": "每日200-500mg维生素C（猕猴桃、柑橘、彩椒），是胶原蛋白合成的必需辅因子，直接影响韧带修复质量。", "level": "重要"},
                {"label": "补钙和维D", "content": "每日1000mg钙（牛奶、豆腐）+ 维生素D（晒太阳15分钟或食用蛋黄、强化食品），维持骨骼和肌肉功能。", "level": "一般"},
                {"label": "水分补充", "content": "每日饮水量不低于2000ml，训练中每15分钟补水200ml。充足水分有助于关节滑液分泌。", "level": "一般"},
                {"label": "避免刺激食物", "content": f"{'恢复期' if stage in ['急性期','亚急性期','恢复期'] else ''}内避免饮酒、减少辛辣油腻食物，以免加重炎症反应。", "level": "注意"},
            ]
        },
        "行": {
            "icon": "🚶",
            "title": "日常行为与姿势管理",
            "tips": [
                {"label": "上下楼梯", "content": "上楼梯以健侧（右腿）先上，下楼梯以患侧（左腿）先下，或侧身借助扶手。避免负重上下楼梯。", "level": "重要"},
                {"label": "久坐管理", "content": "避免膝关节长时间屈曲超过90°（如深蹲式坐姿）。每坐45分钟站立活动5分钟，防止膝关节僵硬。", "level": "重要"},
                {"label": "训练后护理", "content": "每次训练后冰敷膝关节10-15分钟，抬高腿部休息。如有明显肿胀，次日减少训练量。", "level": "重要"},
                {"label": "睡姿调整", "content": "侧卧时双膝间夹枕头，仰卧时膝下垫薄枕（约15°），减少夜间膝关节压力和晨僵感。", "level": "一般"},
                {"label": "负重搬运", "content": "避免单侧搬运重物，必须搬运时双侧均等分配。负重跑步（如背包跑）须待强化期后才可进行。", "level": "注意"},
                {"label": "跪姿避免", "content": f"{'恢复期前' if stage in ['急性期','亚急性期','恢复期'] else ''}避免长时间跪姿（如整理装备时跪地），改用蹲姿或坐姿代替。", "level": "注意"},
            ]
        }
    }

    shoulder_tips = {
        "衣": {
            "icon": "👕", "title": "服装与防护装备",
            "tips": [
                {"label": "护肩使用", "content": "急性期和亚急性期可使用肩部支具或吊带减少肩关节负重，恢复期后改用弹性护肩。", "level": "重要"},
                {"label": "背包方式", "content": "双肩包保持双肩均衡负重，禁止单侧斜跨重包。包重不超过体重的15%。", "level": "重要"},
                {"label": "睡衣枕头", "content": "睡眠时在患侧腋下垫枕头，维持肩关节自然外展15°，避免肩关节受压。", "level": "一般"},
            ]
        },
        "食": {
            "icon": "🥗", "title": "营养与饮食建议",
            "tips": [
                {"label": "蛋白质优先", "content": "每日蛋白质1.6-2.0g/kg，肩袖修复需要充足氨基酸，优先选择乳清蛋白、瘦肉、鱼类。", "level": "重要"},
                {"label": "维生素C+E", "content": "维生素C（胶原合成）+ 维生素E（抗氧化），组合使用对肌腱修复有协同效果。", "level": "重要"},
                {"label": "姜黄素", "content": "天然抗炎成分，可在餐食中加入姜黄粉，或配合黑胡椒提高吸收率。", "level": "一般"},
            ]
        },
        "行": {
            "icon": "🚶", "title": "日常行为与姿势管理",
            "tips": [
                {"label": "举手动作", "content": "避免反复过头顶动作（如晾衣服、拿高处物品），必要时使用工具辅助。", "level": "重要"},
                {"label": "姿势管理", "content": "保持肩膀自然下沉，避免长期耸肩。工作站调整至肘部与桌面等高，减少肩部代偿。", "level": "重要"},
                {"label": "睡眠姿势", "content": "避免患侧卧，选择健侧卧或仰卧，患肩下垫薄枕支撑。", "level": "注意"},
            ]
        }
    }

    back_tips = {
        "衣": {
            "icon": "👕", "title": "服装与防护装备",
            "tips": [
                {"label": "腰部支撑", "content": "急性期和亚急性期使用腰部支撑带，恢复期后逐步减少依赖，过度依赖会导致核心肌群弱化。", "level": "重要"},
                {"label": "鞋具减震", "content": "选择有良好缓震性能的运动鞋，减少行走和跑步时的脊柱冲击力。", "level": "一般"},
            ]
        },
        "食": {
            "icon": "🥗", "title": "营养与饮食建议",
            "tips": [
                {"label": "钙和维D", "content": "重点补充钙（每日1200mg）和维生素D（每日800IU），维持椎骨和椎间盘健康。", "level": "重要"},
                {"label": "胶原蛋白", "content": "猪蹄、骨头汤含丰富胶原前体，配合维C有助于椎间盘胶原合成。", "level": "一般"},
                {"label": "控制体重", "content": "每增加1kg体重，腰椎承受约5kg额外压力。保持健康体重是腰椎保护的长期策略。", "level": "注意"},
            ]
        },
        "行": {
            "icon": "🚶", "title": "日常行为与姿势管理",
            "tips": [
                {"label": "弯腰姿势", "content": "捡物时蹲下（髋膝弯曲），不弯腰驼背。搬运重物时物体贴近身体，避免远离重心搬运。", "level": "重要"},
                {"label": "久坐管理", "content": "坐姿保持腰椎自然弯曲（腰垫辅助），每30分钟站立活动，避免超过45分钟连续坐姿。", "level": "重要"},
                {"label": "睡眠姿势", "content": "侧卧双膝间夹枕（减少腰椎旋转）或仰卧膝下垫枕（减少腰椎过伸）均可。趴睡对腰椎有害。", "level": "注意"},
            ]
        }
    }

    # 按部位选择提示
    tips_map = {
        "膝关节": knee_tips,
        "肩关节": shoulder_tips,
        "腰部": back_tips,
    }
    return tips_map.get(region, knee_tips)


# ─── 切换当前 session 患者（供测试/演示用）────────────────────────────────────

@app.route("/api/load-patient/<int:pid>", methods=["POST"])
def load_patient(pid):
    """将 session 绑定到指定患者 ID（用于演示数据加载）"""
    patient = db.get_patient(pid)
    if not patient:
        return jsonify({"error": "患者不存在"}), 404
    session_state["patient_id"] = pid
    session_state["patient"] = {
        "name": patient.get("name", ""),
        "gender": patient.get("gender", ""),
        "age": patient.get("age"),
        "height": patient.get("height"),
        "weight": patient.get("weight"),
        "injury_description": patient.get("injury_description", ""),
        "recovery_stage": patient.get("recovery_stage", "未知"),
        "pain_level": patient.get("pain_level"),
        "notes": patient.get("notes", ""),
    }
    return jsonify({"success": True, "name": patient.get("name"), "patient_id": pid})


# ─── 完整训练计划 API ─────────────────────────────────────────────────────────

@app.route("/api/training-plan", methods=["GET"])
def get_training_plan():
    """获取当前患者的完整分周期训练计划"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify(None)
    patient = db.get_patient(patient_id)
    if not patient:
        return jsonify(None)
    plan = generate_full_training_plan(patient)
    return jsonify(plan)


# ─── 初版方案重新生成 ─────────────────────────────────────────────────────────

@app.route("/api/plan/generate", methods=["POST"])
def generate_plan():
    """根据当前档案重新生成初版训练方案（无需AI）"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "请先建立个人档案"}), 400
    patient = db.get_patient(patient_id)
    if not patient:
        return jsonify({"error": "患者不存在"}), 404
    auto_plan = generate_default_plan(patient)
    sid = session_state.get("session_id") or None
    db.save_rehab_plan(sid, patient_id, auto_plan)
    return jsonify({"success": True, "message": "初版训练方案已生成", "plan": auto_plan})


# ─── 打卡 API ─────────────────────────────────────────────────────────────────

@app.route("/api/plan/current", methods=["GET"])
def get_current_plan():
    """获取当前患者最新康复方案"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify(None)
    plan = db.get_latest_plan(patient_id)
    return jsonify(plan)


@app.route("/api/checkin", methods=["POST"])
def submit_checkin():
    """提交训练打卡"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "请先建立档案或与助手对话"}), 400

    data = request.get_json() or {}
    if not data.get("exercises"):
        return jsonify({"error": "训练记录不能为空"}), 400

    checkin_id = db.create_checkin(patient_id, data)
    return jsonify({"success": True, "checkin_id": checkin_id, "message": "打卡记录已保存"})


@app.route("/api/checkins", methods=["GET"])
def get_checkins():
    """获取当前患者的打卡历史"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify([])
    checkins = db.get_patient_checkins(patient_id)
    return jsonify(checkins)


# ─── 模型列表 & 对话 & 重置 ──────────────────────────────────────────────────

@app.route("/models", methods=["GET"])
def get_models():
    result = {}
    for pid, pdata in PROVIDERS.items():
        result[pid] = {
            "name": pdata["name"],
            "models": pdata["models"],
            "key_hint": pdata.get("key_hint", "API Key"),
            "key_placeholder": pdata.get("key_placeholder", "sk-..."),
        }
    return jsonify(result)


@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json()
    user_message = data.get("message", "").strip()
    api_key = data.get("api_key", "").strip()
    provider = data.get("provider", "claude").strip()
    model = data.get("model", "claude-sonnet-4-6").strip()

    if not user_message:
        return jsonify({"error": "消息不能为空"}), 400
    if not api_key:
        return jsonify({"error": "请输入 API Key"}), 400

    try:
        result = run_agent_turn(user_message, api_key, provider, model)
        return jsonify(result)
    except anthropic.AuthenticationError:
        return jsonify({"error": "API Key 无效，请检查后重试"}), 401
    except anthropic.RateLimitError:
        return jsonify({"error": "请求频率超限，请稍后重试"}), 429
    except Exception as e:
        try:
            import openai
            if isinstance(e, openai.AuthenticationError):
                return jsonify({"error": "API Key 无效，请检查后重试"}), 401
            if isinstance(e, openai.RateLimitError):
                return jsonify({"error": "请求频率超限，请稍后重试"}), 429
        except ImportError:
            pass
        return jsonify({"error": f"请求失败：{str(e)}"}), 500


@app.route("/reset", methods=["POST"])
def reset():
    db.end_session(session_state.get("session_id"))
    session_state["patient"] = None
    session_state["messages"] = []
    session_state["rehab_plan"] = None
    session_state["feedback_log"] = []
    session_state["session_id"] = None
    session_state["patient_id"] = None
    return jsonify({"success": True, "message": "会话已重置"})


@app.route("/state", methods=["GET"])
def get_state():
    return jsonify({
        "patient": session_state["patient"],
        "rehab_plan": session_state["rehab_plan"],
        "message_count": len(session_state["messages"])
    })


# ─── 教员端 API ───────────────────────────────────────────────────────────────

@app.route("/teacher/api/reviews", methods=["POST"])
def teacher_create_review():
    """教员提交复核/批注"""
    data = request.get_json() or {}
    if not data.get("patient_id"):
        return jsonify({"error": "patient_id 不能为空"}), 400
    if not data.get("content"):
        return jsonify({"error": "复核内容不能为空"}), 400

    review_id = db.create_teacher_review(data)
    return jsonify({"success": True, "review_id": review_id, "message": "复核已提交"})


@app.route("/teacher/api/reviews/<int:patient_id>", methods=["GET"])
def teacher_get_reviews(patient_id):
    """获取某学员的所有教员复核记录"""
    reviews = db.get_patient_reviews(patient_id)
    return jsonify(reviews)


@app.route("/teacher/api/reviews/<int:review_id>/status", methods=["PUT"])
def teacher_update_review_status(review_id):
    """更新复核状态"""
    data = request.get_json() or {}
    status = data.get("status", "")
    if status not in ["待处理", "已确认", "需调整"]:
        return jsonify({"error": "无效状态"}), 400
    db.update_review_status(review_id, status)
    return jsonify({"message": "状态已更新"})


# ─── 智能学伴 对话 ────────────────────────────────────────────────────────────

_COMPANION_SYSTEM = """你是一名专业的运动康复智能助手，服务于正在康复训练中的用户。

你的职责：
1. 用亲切、专业的语言解答用户关于康复训练的疑问
2. 结合用户当前的伤情信息、训练方案和打卡记录，给出个性化建议
3. 当用户提出需要调整训练计划时，或者你判断当前方案需要优化时，在回复末尾给出结构化的方案更新建议

【输出格式规范】
- 正常回复用自然语言，亲切易懂
- 如需更新训练计划，在回复最后附加如下格式（不要附加到正常对话中）：

<<PLAN_UPDATE>>
{
  "summary": "本次调整的一句话概述",
  "reason": "调整原因",
  "new_goal": "（可选）新的训练目标",
  "new_actions": [
    {"名称": "动作名", "说明": "动作说明", "次数": "次数组数", "休息": "休息时间", "注意": "注意事项"}
  ],
  "keep_actions": ["保留的动作名1", "保留的动作名2"],
  "stage_change": "（可选）建议切换到的阶段名称，如强化期"
}
<<END_PLAN_UPDATE>>

【注意】
- 只在确实需要调整方案时才输出 PLAN_UPDATE，普通问答不要输出
- 所有回复使用中文
- 不要过度建议就医，专注于实际的训练指导
"""

def _build_companion_context(patient: dict, plan: dict | None, checkins: list) -> str:
    """构建患者上下文信息供 AI 参考"""
    ctx = f"""【患者信息】
姓名：{patient.get('name','')}，{patient.get('gender','')}，{patient.get('age','')}岁
伤情：{patient.get('injury_description','')}
康复阶段：{patient.get('recovery_stage','')}
当前疼痛评分：{patient.get('pain_level','')} / 10
既往病史：{patient.get('medical_history','无')}
"""
    if plan:
        p = plan.get("plan", {})
        ctx += f"""
【当前训练方案】
目标：{p.get('目标', '')}
训练原则：{p.get('原则', '')}
训练动作：{', '.join(a.get('名称','') for a in p.get('动作', []))}
禁忌：{'; '.join(p.get('禁忌', []))}
"""
    if checkins:
        recent = checkins[:5]
        ctx += "\n【近期打卡记录（最近5次）】\n"
        for c in recent:
            ctx += f"- {c.get('checkin_date','')[:10]}：感受={c.get('overall_feeling','')}, 疼痛={c.get('pain_level','')}/10, 完成率={c.get('completion_rate',0)}%"
            if c.get('symptoms'):
                ctx += f"，症状：{c['symptoms']}"
            ctx += "\n"
    return ctx


@app.route("/api/companion/chat", methods=["POST"])
def companion_chat():
    """智能学伴对话接口"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "请先建立个人档案"}), 400

    data = request.get_json() or {}
    user_msg = data.get("message", "").strip()
    if not user_msg:
        return jsonify({"error": "消息不能为空"}), 400

    # 读取系统配置的 AI 参数
    cfg = db.get_system_config()
    provider = cfg.get("ai_provider", "claude")
    model    = cfg.get("ai_model", "claude-sonnet-4-6")
    api_key  = cfg.get("ai_api_key", "")
    base_url = cfg.get("ai_base_url", "").strip() or None
    if not api_key:
        return jsonify({"error": "管理员尚未配置 AI API Key，请联系管理员在管理端完成配置"}), 400

    # 获取患者上下文
    patient  = db.get_patient(patient_id) or {}
    plan_rec = db.get_latest_plan(patient_id)
    checkins = db.get_patient_checkins(patient_id, limit=10)

    ctx = _build_companion_context(patient, plan_rec, checkins)

    # RAG：从知识库检索与用户问题相关的内容片段
    rag_context = ""
    try:
        kb_hits = db.kb_search(user_msg, top_k=4)
        if kb_hits:
            rag_context = "\n\n## 知识库参考内容\n以下内容来自知识库，供回答参考（按相关性排序）：\n"
            for hit in kb_hits:
                rag_context += f"\n【{hit['item_title']}】\n{hit['content']}\n"
    except Exception:
        pass  # 知识库检索失败不影响正常对话

    system_prompt = _COMPANION_SYSTEM + "\n\n" + ctx + rag_context

    # 构建消息历史（从 data 传入历史，最多保留10轮）
    history = data.get("history", [])[-20:]
    messages = history + [{"role": "user", "content": user_msg}]

    try:
        final_text, _ = call_llm(
            provider, model, api_key,
            messages, [],          # 学伴不使用 tools
            system_prompt,
            lambda n, i: {},       # 空的 tool processor
            base_url=base_url,
        )
    except Exception as e:
        return jsonify({"error": f"AI 调用失败：{str(e)}"}), 500

    # 解析 PLAN_UPDATE 块
    plan_suggestion = None
    reply = final_text
    if "<<PLAN_UPDATE>>" in final_text and "<<END_PLAN_UPDATE>>" in final_text:
        start = final_text.index("<<PLAN_UPDATE>>") + len("<<PLAN_UPDATE>>")
        end   = final_text.index("<<END_PLAN_UPDATE>>")
        json_str = final_text[start:end].strip()
        reply    = final_text[:final_text.index("<<PLAN_UPDATE>>")].strip()
        try:
            plan_suggestion = json.loads(json_str)
        except Exception:
            plan_suggestion = {"summary": "AI 建议调整训练计划", "raw": json_str}

    return jsonify({
        "reply": reply,
        "plan_suggestion": plan_suggestion,
    })


@app.route("/api/companion/apply-plan", methods=["POST"])
def companion_apply_plan():
    """将学伴建议的方案更新应用到训练计划"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "未找到患者信息"}), 400

    data       = request.get_json() or {}
    suggestion = data.get("plan_suggestion", {})
    if not suggestion:
        return jsonify({"error": "无方案数据"}), 400

    # 获取当前方案，在其基础上更新
    existing  = db.get_latest_plan(patient_id)
    old_plan  = (existing or {}).get("plan", {})

    new_actions = suggestion.get("new_actions", [])
    keep_names  = suggestion.get("keep_actions", [])
    old_actions = old_plan.get("动作", [])

    # 保留指定动作 + 添加新动作
    if new_actions or keep_names:
        kept = [a for a in old_actions if a.get("名称","") in keep_names] if keep_names else old_actions
        merged = kept + [a for a in new_actions if a.get("名称","") not in [x.get("名称","") for x in kept]]
    else:
        merged = old_actions

    new_stage = suggestion.get("stage_change") or old_plan.get("stage") or existing.get("stage","")
    new_goal  = suggestion.get("new_goal") or old_plan.get("目标","")

    updated_plan = {**old_plan,
        "目标":    new_goal,
        "动作":    merged,
        "stage":   new_stage,
        "is_auto_generated": False,   # 标记为 AI 细化版本
        "updated_by": "智能学伴",
    }
    if suggestion.get("stage_change"):
        updated_plan["stage"] = suggestion["stage_change"]

    sid = session_state.get("session_id") or None
    db.save_rehab_plan(sid, patient_id, updated_plan)

    return jsonify({"success": True, "message": "训练计划已更新"})


# ═══════════════════════════════════════════════════════════════════════════════
# 管理端路由
# ═══════════════════════════════════════════════════════════════════════════════

@app.route("/admin/api/stats", methods=["GET"])
def admin_stats():
    return jsonify(db.get_stats())


# ─── 患者管理 ─────────────────────────────────────────────────────────────────

@app.route("/admin/api/patients", methods=["GET"])
def admin_patients():
    search = request.args.get("search", "")
    stage = request.args.get("stage", "")
    role = request.args.get("role", "")
    page = int(request.args.get("page", 1))
    return jsonify(db.get_patients(search, stage, role, page))


@app.route("/admin/api/patients", methods=["POST"])
def admin_create_patient():
    data = request.get_json() or {}
    if not data.get("name"):
        return jsonify({"error": "姓名不能为空"}), 400
    pid = db.create_patient(data)
    return jsonify({"id": pid, "message": "患者创建成功"})


@app.route("/admin/api/patients/<int:patient_id>", methods=["GET"])
def admin_get_patient(patient_id):
    p = db.get_patient(patient_id)
    if not p:
        return jsonify({"error": "患者不存在"}), 404
    return jsonify(p)


@app.route("/admin/api/patients/<int:patient_id>", methods=["PUT"])
def admin_update_patient(patient_id):
    data = request.get_json() or {}
    db.update_patient(patient_id, data)
    return jsonify({"message": "更新成功"})


@app.route("/admin/api/patients/<int:patient_id>", methods=["DELETE"])
def admin_delete_patient(patient_id):
    db.delete_patient(patient_id)
    return jsonify({"message": "删除成功"})


@app.route("/admin/api/patients/<int:patient_id>/generate-profile", methods=["POST"])
def admin_generate_profile(patient_id):
    data = request.get_json() or {}
    text = data.get("text", "")
    api_key = data.get("api_key", "")
    provider = data.get("provider", "claude")
    model = data.get("model", "claude-sonnet-4-6")

    if not api_key:
        return jsonify({"error": "请提供 API Key"}), 400

    patient = db.get_patient(patient_id)
    context = ""
    if patient:
        context = f"患者姓名：{patient.get('name', '')}，伤情：{patient.get('injury_description', '')}，阶段：{patient.get('recovery_stage', '')}。"

    prompt = f"""请根据以下康复相关信息，为患者生成一份简洁的综合康复画像（200字以内），包括：伤情特点、当前阶段、康复重点、注意事项。

{context}

补充信息：
{text}

请直接输出画像文字，不要加标题或额外说明。"""

    try:
        if provider == "claude":
            client = anthropic.Anthropic(api_key=api_key)
            resp = client.messages.create(
                model=model,
                max_tokens=400,
                messages=[{"role": "user", "content": prompt}]
            )
            profile_text = resp.content[0].text
        else:
            import openai
            pdata = PROVIDERS.get(provider, {})
            base_url = pdata.get("base_url", "")
            client = openai.OpenAI(api_key=api_key, base_url=base_url if base_url else None)
            resp = client.chat.completions.create(
                model=model,
                max_tokens=400,
                messages=[{"role": "user", "content": prompt}]
            )
            profile_text = resp.choices[0].message.content

        db.update_patient(patient_id, {"ai_profile": profile_text})
        return jsonify({"profile": profile_text, "message": "画像生成成功"})

    except Exception as e:
        return jsonify({"error": f"生成失败：{str(e)}"}), 500


# ─── 会话日志 ─────────────────────────────────────────────────────────────────

@app.route("/admin/api/sessions", methods=["GET"])
def admin_sessions():
    page = int(request.args.get("page", 1))
    return jsonify(db.get_sessions(page))


@app.route("/admin/api/tool-calls", methods=["GET"])
def admin_tool_calls():
    page = int(request.args.get("page", 1))
    return jsonify(db.get_tool_calls(page))


@app.route("/admin/api/feedback", methods=["GET"])
def admin_feedback():
    page = int(request.args.get("page", 1))
    risk = request.args.get("risk", "")
    return jsonify(db.get_feedback_logs(page, risk=risk))


# ─── 动作库 ───────────────────────────────────────────────────────────────────

@app.route("/admin/api/actions", methods=["GET"])
def admin_actions():
    search = request.args.get("search", "")
    stage = request.args.get("stage", "")
    page = int(request.args.get("page", 1))
    return jsonify(db.get_actions(search, stage, page))


@app.route("/admin/api/actions", methods=["POST"])
def admin_create_action():
    data = request.get_json() or {}
    if not data.get("name"):
        return jsonify({"error": "动作名称不能为空"}), 400
    aid = db.create_action(data)
    return jsonify({"id": aid, "message": "动作创建成功"})


@app.route("/admin/api/actions/<int:action_id>", methods=["GET"])
def admin_get_action(action_id):
    a = db.get_action(action_id)
    if not a:
        return jsonify({"error": "动作不存在"}), 404
    return jsonify(a)


@app.route("/admin/api/actions/<int:action_id>", methods=["PUT"])
def admin_update_action(action_id):
    data = request.get_json() or {}
    db.update_action(action_id, data)
    return jsonify({"message": "更新成功"})


@app.route("/admin/api/actions/<int:action_id>", methods=["DELETE"])
def admin_delete_action(action_id):
    db.delete_action(action_id)
    return jsonify({"message": "删除成功"})


# ─── 系统用户 ─────────────────────────────────────────────────────────────────

@app.route("/admin/api/users", methods=["GET"])
def admin_users():
    return jsonify(db.get_system_users())


@app.route("/admin/api/users", methods=["POST"])
def admin_create_user():
    data = request.get_json() or {}
    if not data.get("username"):
        return jsonify({"error": "用户名不能为空"}), 400
    try:
        uid = db.create_system_user(data)
        return jsonify({"id": uid, "message": "用户创建成功"})
    except Exception as e:
        return jsonify({"error": f"创建失败：{str(e)}"}), 400


@app.route("/admin/api/users/<int:user_id>", methods=["PUT"])
def admin_update_user(user_id):
    data = request.get_json() or {}
    db.update_system_user(user_id, data)
    return jsonify({"message": "更新成功"})


@app.route("/admin/api/users/<int:user_id>", methods=["DELETE"])
def admin_delete_user(user_id):
    db.delete_system_user(user_id)
    return jsonify({"message": "删除成功"})


# ─── 管理端 AI 配置 ───────────────────────────────────────────────────────────

@app.route("/admin/api/ai-config", methods=["GET"])
def admin_get_ai_config():
    cfg = db.get_system_config()
    return jsonify({
        "ai_provider": cfg.get("ai_provider", "claude"),
        "ai_model":    cfg.get("ai_model", "claude-sonnet-4-6"),
        "ai_api_key":  cfg.get("ai_api_key", ""),
        "ai_base_url": cfg.get("ai_base_url", ""),
        "ai_system_prompt_extra": cfg.get("ai_system_prompt_extra", ""),
    })


@app.route("/admin/api/ai-config", methods=["POST"])
def admin_save_ai_config():
    data = request.get_json() or {}
    allowed = {"ai_provider", "ai_model", "ai_api_key", "ai_base_url", "ai_system_prompt_extra"}
    save_data = {k: v for k, v in data.items() if k in allowed}
    if not save_data:
        return jsonify({"error": "无有效配置项"}), 400
    db.set_system_config_batch(save_data)
    return jsonify({"success": True, "message": "AI 配置已保存"})


@app.route("/admin/api/ai-config/test", methods=["POST"])
def admin_test_ai_connection():
    """测试当前 AI 配置是否可用"""
    cfg = db.get_system_config()
    provider = cfg.get("ai_provider", "claude")
    model = cfg.get("ai_model", "claude-sonnet-4-6")
    api_key = cfg.get("ai_api_key", "")
    saved_base_url = cfg.get("ai_base_url", "").strip() or None
    if not api_key:
        return jsonify({"error": "尚未配置 API Key，请先保存配置"}), 400
    try:
        test_messages = [
            {"role": "system", "content": "你是测试助手"},
            {"role": "user", "content": "请回复：连接正常"},
        ]
        if provider == "claude" and not saved_base_url:
            import anthropic as anthropic_sdk
            client = anthropic_sdk.Anthropic(api_key=api_key)
            resp = client.messages.create(
                model=model, max_tokens=64,
                messages=[{"role": "user", "content": "请回复：连接正常"}]
            )
            reply = resp.content[0].text if resp.content else "ok"
        else:
            base_url = saved_base_url or PROVIDERS.get(provider, {}).get("base_url", "https://api.openai.com/v1")
            data = llm_providers._openai_chat(base_url, api_key, model, test_messages, timeout=30)
            reply = data["choices"][0]["message"].get("content") or "ok"
        return jsonify({"success": True, "message": f"模型响应正常，内容：{reply[:50]}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ─── 管理端 AI 配置（获取可用模型列表供选择）─────────────────────────────────

@app.route("/admin/api/ai-models", methods=["GET"])
def admin_get_ai_models():
    result = {}
    for pid, pdata in PROVIDERS.items():
        result[pid] = {
            "name": pdata["name"],
            "models": pdata["models"],
            "key_placeholder": pdata.get("key_placeholder", "sk-..."),
        }
    return jsonify(result)


# ─── 病历档案上传与解析 ────────────────────────────────────────────────────────

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads", "medical_records")
ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".png", ".jpg", ".jpeg"}
MAX_FILE_MB = 20


def _parse_pdf(path: str) -> str:
    import pdfplumber
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return "\n\n".join(texts)


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _parse_txt(path: str) -> str:
    import chardet
    with open(path, "rb") as f:
        raw = f.read()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    return raw.decode(enc, errors="replace")


def _parse_xlsx(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        rows.append(f"【表格：{ws.title}】")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
    return "\n".join(rows)


def _parse_image(path: str) -> str:
    """图片暂不 OCR，返回提示"""
    return f"[图片文件已上传，暂不支持自动识别内容。如需解读，请将图片内容手动填写至伤情描述中。]"


def parse_medical_file(path: str, ext: str) -> str:
    """根据扩展名分派解析器，返回提取的文本"""
    try:
        if ext == ".pdf":
            return _parse_pdf(path)
        elif ext == ".docx":
            return _parse_docx(path)
        elif ext in (".txt", ".md"):
            return _parse_txt(path)
        elif ext == ".xlsx":
            return _parse_xlsx(path)
        elif ext in (".png", ".jpg", ".jpeg"):
            return _parse_image(path)
        else:
            return _parse_txt(path)   # fallback
    except Exception as e:
        return f"[文件解析出错：{e}]"


@app.route("/api/upload-medical-record", methods=["POST"])
def upload_medical_record():
    """上传病历档案，解析文本内容"""
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "文件名为空"}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ALLOWED_EXT:
        return jsonify({"error": f"不支持的文件格式 {ext}，支持：PDF / DOCX / TXT / XLSX / 图片"}), 400

    file_size = len(f.read())
    f.seek(0)
    if file_size > MAX_FILE_MB * 1024 * 1024:
        return jsonify({"error": f"文件过大，最大支持 {MAX_FILE_MB}MB"}), 400

    # 保存文件
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(UPLOAD_DIR, safe_name)
    f.save(save_path)

    # 解析文本
    text = parse_medical_file(save_path, ext)
    text = text.strip()

    # 截断超长内容（避免前端卡死）
    truncated = len(text) > 8000
    if truncated:
        text = text[:8000] + "\n\n...[内容过长，已截断显示前8000字]"

    return jsonify({
        "success": True,
        "filename": f.filename,
        "ext": ext,
        "text": text,
        "truncated": truncated,
        "char_count": len(text),
        "saved_as": safe_name,
    })


@app.route("/api/medical-records", methods=["GET"])
def list_medical_records():
    """列出当前患者已上传的病历档案（元数据）"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify([])
    cfg = db.get_system_config()
    records_json = cfg.get(f"medical_records_{patient_id}", "[]")
    try:
        records = json.loads(records_json)
    except Exception:
        records = []
    return jsonify(records)


@app.route("/api/medical-records", methods=["POST"])
def save_medical_record_meta():
    """保存病历档案元数据（文件名、摘要等）到 system_config"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "未绑定患者"}), 400
    data = request.get_json() or {}
    cfg = db.get_system_config()
    key = f"medical_records_{patient_id}"
    try:
        records = json.loads(cfg.get(key, "[]"))
    except Exception:
        records = []
    records.append({
        "id": uuid.uuid4().hex[:8],
        "filename": data.get("filename", ""),
        "saved_as": data.get("saved_as", ""),
        "ext": data.get("ext", ""),
        "char_count": data.get("char_count", 0),
        "uploaded_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "summary": data.get("summary", ""),   # 用户填写的摘要/用途说明
    })
    db.set_system_config(key, json.dumps(records, ensure_ascii=False))
    return jsonify({"success": True, "count": len(records)})


@app.route("/api/medical-records/<record_id>", methods=["DELETE"])
def delete_medical_record(record_id):
    """删除某条病历档案元数据"""
    patient_id = session_state.get("patient_id")
    if not patient_id:
        return jsonify({"error": "未绑定患者"}), 400
    cfg = db.get_system_config()
    key = f"medical_records_{patient_id}"
    try:
        records = json.loads(cfg.get(key, "[]"))
    except Exception:
        records = []
    to_delete = next((r for r in records if r["id"] == record_id), None)
    if to_delete:
        # 删除磁盘文件
        disk_path = os.path.join(UPLOAD_DIR, to_delete.get("saved_as", ""))
        if os.path.exists(disk_path):
            os.remove(disk_path)
        records = [r for r in records if r["id"] != record_id]
        db.set_system_config(key, json.dumps(records, ensure_ascii=False))
    return jsonify({"success": True})


# ─── 知识库路由 ────────────────────────────────────────────────────────────────

KB_UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads", "knowledge")
KB_ALLOWED_EXT = {".pdf", ".docx", ".doc", ".txt", ".md",
                  ".epub", ".mobi", ".xlsx",
                  ".mp4", ".mov", ".avi", ".mkv", ".srt"}
KB_MAX_MB = 200
KB_CHUNK_SIZE = 400    # 每块字符数
KB_CHUNK_OVERLAP = 80  # 相邻块重叠字符数


def _kb_chunk_text(text: str) -> list[str]:
    """将长文本分割为带重叠的块"""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + KB_CHUNK_SIZE
        chunks.append(text[start:end])
        start += KB_CHUNK_SIZE - KB_CHUNK_OVERLAP
    return [c.strip() for c in chunks if c.strip()]


def _kb_extract_text(path: str, ext: str) -> str:
    """从各类文件提取纯文本"""
    try:
        if ext == ".pdf":
            return _parse_pdf(path)
        elif ext in (".docx", ".doc"):
            return _parse_docx(path)
        elif ext in (".txt", ".md", ".srt"):
            return _parse_txt(path)
        elif ext == ".xlsx":
            return _parse_xlsx(path)
        elif ext == ".epub":
            import ebooklib
            from ebooklib import epub
            from html.parser import HTMLParser

            class _TagStripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.parts = []
                def handle_data(self, data):
                    self.parts.append(data)
                def get_text(self):
                    return " ".join(self.parts)

            book = epub.read_epub(path)
            parts = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                ts = _TagStripper()
                ts.feed(item.get_content().decode("utf-8", errors="replace"))
                parts.append(ts.get_text())
            return "\n\n".join(parts)
        elif ext == ".mobi":
            # mobi 用 pdfplumber 无法读；尝试作为 epub 读（部分 mobi 可行）
            try:
                import ebooklib
                from ebooklib import epub
                book = epub.read_epub(path)
                parts = []
                for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                    parts.append(item.get_content().decode("utf-8", errors="replace"))
                return "\n\n".join(parts)
            except Exception:
                return "[MOBI 文件已上传，当前版本暂不支持自动提取文本内容。如需检索，请将关键内容粘贴为 TXT 格式上传。]"
        elif ext in (".mp4", ".mov", ".avi", ".mkv"):
            return f"[视频文件：{os.path.basename(path)}。视频内容不自动转录；若有对应字幕文件（.srt），请单独上传以支持检索。]"
        else:
            return _parse_txt(path)
    except Exception as e:
        return f"[文件解析出错：{e}]"


@app.route("/api/knowledge/upload", methods=["POST"])
def kb_upload():
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400
    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "文件名为空"}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in KB_ALLOWED_EXT:
        return jsonify({"error": f"不支持的格式 {ext}"}), 400

    raw = f.read()
    if len(raw) > KB_MAX_MB * 1024 * 1024:
        return jsonify({"error": f"文件超过 {KB_MAX_MB}MB 限制"}), 400

    os.makedirs(KB_UPLOAD_DIR, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(KB_UPLOAD_DIR, safe_name)
    with open(save_path, "wb") as fout:
        fout.write(raw)

    # 提取文本 & 分块
    text = _kb_extract_text(save_path, ext)
    chunks = _kb_chunk_text(text) if text and not text.startswith("[") else []

    title    = request.form.get("title") or os.path.splitext(f.filename)[0]
    category = request.form.get("category", "")
    tags     = request.form.get("tags", category)  # category stored as tags
    desc     = request.form.get("description", "")

    item_id = db.kb_add_item(
        title=title, filename=f.filename, file_type=ext.lstrip("."),
        file_size=len(raw), file_path=save_path,
        content_text=text[:2000],   # 存摘要（前2000字）
        tags=tags, description=desc,
    )
    if chunks:
        db.kb_add_chunks(item_id, chunks)

    return jsonify({
        "success": True, "id": item_id, "title": title,
        "chunk_count": len(chunks), "char_count": len(text),
        "preview": text[:300] if text else "",
    })


@app.route("/api/knowledge/items", methods=["GET"])
def kb_list():
    items = db.kb_list_items()
    for item in items:
        item["file_size_kb"] = round(item.get("file_size", 0) / 1024, 1)
        item["category"] = item.get("tags", "")
    return jsonify({"items": items})


@app.route("/api/knowledge/items/<int:item_id>", methods=["GET"])
def kb_get(item_id):
    item = db.kb_get_item(item_id)
    if not item:
        return jsonify({"error": "未找到"}), 404
    return jsonify({
        "id": item["id"],
        "title": item["title"],
        "category": item.get("tags", ""),
        "preview": item.get("content_text", "") or "（无可预览的文本内容）",
    })


@app.route("/api/knowledge/items/<int:item_id>", methods=["DELETE"])
def kb_delete(item_id):
    file_path = db.kb_delete_item(item_id)
    if file_path and os.path.exists(file_path):
        os.remove(file_path)
    return jsonify({"success": True})


@app.route("/api/knowledge/search", methods=["POST"])
def kb_search():
    data  = request.get_json() or {}
    query = data.get("query", "").strip()
    top_k = min(int(data.get("top_k", 5)), 20)
    if not query:
        return jsonify({"error": "查询不能为空"}), 400
    try:
        results = db.kb_search(query, top_k=top_k)
        return jsonify({"results": results, "query": query, "total": len(results)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    debug = os.environ.get("FLASK_ENV") != "production"
    print("=" * 50)
    print("  训练伤康复智能助手 v2")
    print(f"  用户端: http://localhost:{port}")
    print(f"  打卡页: http://localhost:{port}/checkin")
    print(f"  教师端: http://localhost:{port}/teacher")
    print(f"  管理端: http://localhost:{port}/admin")
    print("=" * 50)
    app.run(host="0.0.0.0", port=port, debug=debug)
